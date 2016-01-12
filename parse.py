# -*- coding: utf-8 -*-

#TODO: fix wrong subjects
#TODO: make circles to links
#TODO: split into smaller HTML files
#TODO: handle footnotes' fonts
#TODO: do something with new-lines?
#TODO: TOC, search
#TODO: make footnotes to be superscript, without using ()

#TODO: remove out 'styles' dict
#TODO: icon
#TODO: automate build
#TODO: iphone?



import docx
import docx_fork_ludoo
import dominate
import dominate.tags as tags
import re
import zipfile


#doc_file_name = 'dict.docx'
doc_file_name = 'dict_short.docx'
#doc_file_name = 'snippet2.docx'

# word_doc = docx.Document(doc_file_name)
# word_doc_footnotes = docx_fork_ludoo.Document(doc_file_name)

#word_doc = docx.Document(doc_file_name)
word_doc =  docx_fork_ludoo.Document(doc_file_name)


# support old and new version of docx
# new is preferred, because, well, it's new...
# old is preferred because I'm using a branch with footnotes support
def run_style_id(run):
    try:
        return run.style.style_id
    except:
        if run.style:
            return run.style
        else:
            return 'DefaultParagraphFont'


styles = {
    's01': 'subject',
    's11': 'sub-subject',
    's02': 'definition',
    's03': 'source',
    'Heading3Char': 'definition',
    '1': 'definition',   #?
    'FootnoteTextChar1': 'definition',   #?
    'HebrewChar': 'definition',   #?
    'DefaultParagraphFont': 'source',    #?

    's15': 'subject_small',
    's17': 'subject_small',
    's1510': 'unknown_small',
    's05': 'definition_small',
    's038': 'source_small',
    's0590': 'source_small',
    's050': 'source_small',

    's149': 'subject_light',
    's16': 'subject_light',
    's14': 'subject_light',
    's168': 'sub-subject_light',
    's048': 'definition_light',
    's12': 'definition_light',
    's04': 'unknown_light',
    's127': 'source_light',

    's02Symbol': 's02Symbol',   # MeUyan

    'FootnoteReference': 'FootnoteReference',
    'EndnoteReference': 'EndnoteReference', #?
}

unknown_list = []


def subject(type, text):
    with tags.span(tags.a(text, href="#%s" % text, id=text)):
        tags.attr(cls=type)

def regular(type, text):
    if type == 'footnote':
        with tags.a("(%s)" % text.strip()):
            tags.attr(cls="ptr")
    else:
        with tags.span(text):
            tags.attr(cls=type)

def is_subject(para, i, next=False):
    type, text = para[i]
    # print "is? ", type, text.strip()
    if 'subject' in type and not re.search(r"\w", text, re.UNICODE):
        print "!", text
    return 'subject' in type and re.search(r"\w", text, re.UNICODE)



def analyze_and_fix(para):
    # unite splitted adjacent similar types
    prev_type, prev_text = None, ""
    new_para = []
    for (type, text) in para:
        if prev_type:
            if type == prev_type or text.strip() in ("", u"°"):
                prev_text += text
            else:
                new_para.append((prev_type, prev_text))
                prev_type, prev_text = type, text
        else:
            prev_type, prev_text = type, text
    new_para.append((prev_type, prev_text))

    para = new_para
    new_para = []

    # fix wrong 'source's
    source_pattern = re.compile(r"(\[.*\])")
    for (type, text) in para:
        if type == 'source':
            small = False
            l = source_pattern.split(text)
            for (chunk) in source_pattern.split(text):
                if source_pattern.match(chunk):
                    if small:
                        new_para.append(('source_small', chunk))
                    else:
                        new_para.append((type, chunk))
                elif chunk.strip() != "":
                    new_para.append(('definition_small', chunk))
                    small = True
                # re.split(r"(\[.*\])", s)
        else:
            new_para.append((type, text))

    # fix
    return new_para

def is_prev_subject(para, i):
    try:
        return is_subject(para, i-2) and para[i-1][1].strip() == "-"
    except:
        return False

def add_to_output(para):
    # we shouldn't accept empty paragraph (?)
    assert len(para) > 0

    # if there is only 1 'subject' item in the paragraph, it's probably a heading
    first_type, first_text = para[0]
    if len(para) == 1 and first_type == 'subject':
        # take the 'text' of the first item
        tags.h3(first_text)
    else:
        for (i, (type, text)) in enumerate(para):
            if type == "new_line":
                tags.br()
            elif is_subject(para, i):
                if not is_prev_subject(para, i):
                    tags.p()
                    #tags.br()
                subject(type, text)
            else:
                regular(type, text)

        tags.br()

def add_footnote_to_output(id, paragraphs):
    text = ""
    for (para) in paragraphs:
        text += para.text
    tags.li(text)


temp_l = []
def bold_type(type):
    if type == 'definition':
        return 'subject_small'
    elif type == 'source':
        return 'sub-subject'
    elif 'subject' in type:
        return type
    else:
        if type not in temp_l:
            print "Unexpected bold!", type
            temp_l.append(type)
        return type

html_doc = dominate.document(title=u"מילון הראיה")
html_doc['dir'] = 'rtl'
with html_doc.head:
    tags.link(rel='stylesheet', href='style.css')
    tags.link(rel='stylesheet', href='html_demos-gh-pages/footnotes.css')
    tags.script(src="html_demos-gh-pages/footnotes.js")

with open('debug.txt', 'w') as debug_file:
    with html_doc:
        for (paragraph) in word_doc.paragraphs:
            if paragraph.text.strip():
                # print "Paragraph:", paragraph.text, "$"
                para = []
                debug_file.write("\n\nNEW_PARA:\n------\n")
                for (run) in paragraph.runs:
                    s = "!%s:%s$" % (styles.get(run_style_id(run), run_style_id(run)), run.text)
                    # print "!%s:%s$" % (styles.get(run.style.style_id, run.style.style_id), run.text)
                    debug_file.write(s.encode('utf8'))
                    type = styles.get(run_style_id(run), "unknown")
                    if run.bold:
                        type = bold_type(type)
                    # add_to_output(type, run.text)
                    para.append((type, run.text))

                    if type == "unknown":
                        if run_style_id(run) not in unknown_list:
                            unknown_list.append(run_style_id(run))
                            print paragraph.text
                            s = "\nMissing: !%s:%s$\n\n" % (run_style_id(run), run.text)
                            print s
                            debug_file.write(s.encode('utf8'))

                    if run.footnote_references:
                        for (note) in run.footnote_references:
                            para.append(('footnote', str(note.id)))

                    para.append(("new_line", ""))


                para = analyze_and_fix(para)
                add_to_output(para)
            else:
                # print paragraph
                tags.p()

        # tags.hr()

        with tags.ol(id="footnotes"):
            for footnote in word_doc.footnotes_part.notes:
                if footnote.id >= 1:
                    add_footnote_to_output(footnote.id, footnote.paragraphs)



if unknown_list:
    print "\n\nMissing:"
    print unknown_list

with open('debug.html', 'w') as debug_file:
    debug_file.write(html_doc.render(inline=False).encode('utf8'))
    print "Created debug.html"

with open('index.html', 'w') as debug_file:
    debug_file.write(html_doc.render(inline=True).encode('utf8'))
    print "Created index.html"

with zipfile.ZipFile('milon.zip', 'w', zipfile.ZIP_DEFLATED) as myzip:
    for (filename) in (
        'config.xml',
        'index.html',
        'style.css',
        'html_demos-gh-pages/footnotes.css',
        'html_demos-gh-pages/footnotes.js',
    ):
        myzip.write(filename)

    print "Created milon.zip"